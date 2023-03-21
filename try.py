import pyodbc
import re
import random
import datetime
import docx

#for row in cursor:
  #  a += row[2]
    #print(a)
#Капуста, морковь, лук, петрушка, укроп, свекла, мясо(500 грамм), Фасоль (200 грамм)
#Cabbage, carrot, onion, parsley, dill, beetroot, meat (500 grams), Beans (200 grams)



#Количество каждого ингредиента в блюде.
Cabbage = int(0)
Carrot = int(0)
Onion = int(0)
Parsley = int(0)
Dill = int(0)
Beetroot = int(0)
Meat = int(0)
Potato = int(0)
Garlic = int(0)
TomatoPaste = int(0)
Vinegar = int(0)
Beans = int(0)

CountElement = int(0)


#Количество блюд
Quantity = int(0)
Counter = int(0)


#Выбор борща (свой или стандартный)
Selecter = int(0)


#Сумма для чека
SummCheck = int(0)
Summ = int(0)


#Авторизация
Autoriz = int(0)

#Роль
Role = int(0) #2 - admin | 1 - user


#Баланс
Balance = int(0)

#Проверка на покупку
Challenge = int(0)

#Выход
ExitYes = int(0)

#Провекра наличия продуктов
IsHere = int(0) #2- нет на складе| 0 - есть, но не столько сколько нужно | 1 - успешная операция


#Проверка таракана
Cockroach = 0 #от 1 до 5
UserSees = 0 #от 1 до 6
cok = 0


#Функция подсчета цены за количества продукта
def SummaCheck(Cost, Count):
    CountElement = int(0)
    Summ=int(0)
    while CountElement<Count:
        CountElement+=1
        Summ+=Cost
    return Summ

#Фугкция получения цены за продукт
def GetPrice(Name):
    cnxn = pyodbc.connect('DRIVER={SQL Server};SERVER=LAPTOP-BIF79HGK\ARTYOMSQL;DATABASE=BorshDatabase;UID=sa;PWD=123')
    cursor = cnxn.cursor()
    cursor.execute(f'SELECT * FROM Ingredients WHERE Ingredient_Name = \'{Name}\'')
    for row in cursor:
        a = int(row[3])
        return a
    cnxn.close()

#Функции установки количества ингридиентов
def CabbageFunct (Name):
    a = int(3)
    while a == 3:
        b = int(CheckProduct("Капуста", Name))
        if (b == 1):
            a = 0
        if (b == 2):
            Name = 0
            a = 0
        if (b==0):
            Name = int(input("Введите другое количество: "))
    price = int(GetPrice("Капуста"))
    summ = SummaCheck(int(price), Name)
    return summ

def CarrotFunct (Name):
    a = int(3)
    while a == 3:
        b = int(CheckProduct("Морковь", Name))
        if (b == 1):
            a = 0
        if (b == 2):
            Name = 0
            a = 0
        if (b==0):
            Name = int(input("Введите другое количество: "))
    price = int(GetPrice("Морковь"))
    summ = SummaCheck(int(price), Name)
    return summ

def UnionFunct (Name):
    a = int(3)
    while a == 3:
        b = int(CheckProduct("Луковица", Name))
        if (b == 1):
            a = 0
        if (b == 2):
            Name = 0
            a = 0
        if (b==0):
            Name = int(input("Введите другое количество: "))
    price = int(GetPrice("Луковица"))
    summ = SummaCheck(int(price), Name)
    return summ

def ParsleyFunct (Name):
    a = int(3)
    while a == 3:
        b = int(CheckProduct("Петрушка", Name))
        if (b == 1):
            a = 0
        if (b == 2):
            Name = 0
            a = 0
        if (b==0):
            Name = int(input("Введите другое количество: "))
    price = int(GetPrice("Петрушка"))
    summ = SummaCheck(int(price), Name)
    return summ

def DillFunct (Name):
    a = int(3)
    while a == 3:
        b = int(CheckProduct("Укроп", Name))
        if (b == 1):
            a = 0
        if (b == 2):
            Name = 0
            a = 0
        if (b==0):
            Name = int(input("Введите другое количество: "))
    price = int(GetPrice("Укроп"))
    summ = SummaCheck(int(price), Name)
    return summ

def BeetrootFunct (Name):
    a = int(3)
    while a == 3:
        b = int(CheckProduct("Свекла", Name))
        if (b == 1):
            a = 0
        if (b == 2):
            Name = 0
            a = 0
        if (b==0):
            Name = int(input("Введите другое количество: "))
    price = int(GetPrice("Свекла"))
    summ = SummaCheck(int(price), Name)
    return summ 

def MeatFunct (Name):
    a = int(3)
    while a == 3:
        b = int(CheckProduct("Мясо", Name))
        if (b == 1):
            a = 0
        if (b == 2):
            Name = 0
            a = 0
        if (b==0):
            Name = int(input("Введите другое количество: "))
    price = int(GetPrice("Мясо"))
    summ = SummaCheck(int(price), Name)
    return summ

def PotatoFunct (Name):
    a = int(3)
    while a == 3:
        b = int(CheckProduct("Картофель", Name))
        if (b == 1):
            a = 0
        if (b == 2):
            Name = 0
            a = 0
        if (b==0):
            Name = int(input("Введите другое количество: "))
    price = int(GetPrice("Картофель"))
    summ = SummaCheck(int(price), Name)
    return summ

def GarlicFunct (Name):
    a = int(3)
    while a == 3:
        b = int(CheckProduct("Чеснок", Name))
        if (b == 1):
            a = 0
        if (b == 2):
            Name = 0
            a = 0
        if (b==0):
            Name = int(input("Введите другое количество: "))
    price = int(GetPrice("Чеснок"))
    summ = SummaCheck(int(price), Name)
    return summ

def TomatoPasteFunct (Name):
    a = int(3)
    while a == 3:
        b = int(CheckProduct("Томатная паста", Name))
        if (b == 1):
            a = 0
        if (b == 2):
            Name = 0
            a = 0
        if (b==0):
            Name = int(input("Введите другое количество: "))
    price = int(GetPrice("Томатная паста"))
    summ = SummaCheck(int(price), Name)
    return summ

def VinegarFunct (Name):
    a = int(3)
    while a == 3:
        b = int(CheckProduct("Уксус", Name))
        if (b == 1):
            a = 0
        if (b == 2):
            Name = 0
            a = 0
        if (b==0):
            Name = int(input("Введите другое количество: "))
    price = int(GetPrice("Уксус"))
    summ = SummaCheck(int(price), Name)
    return summ

def BeansFunct (Name):
    a = int(3)
    while a == 3:
        b = int(CheckProduct("Фасоль", Name))
        if (b == 1):
            a = 0
        if (b == 2):
            Name = 0
            a = 0
        if (b==0):
            Name = int(input("Введите другое количество: "))
    price = int(GetPrice("Фасоль"))
    summ = SummaCheck(int(price), Name)
    return summ



#Функция авторизации
def Autorization(Number):
    cnxn = pyodbc.connect('DRIVER={SQL Server};SERVER=LAPTOP-BIF79HGK\ARTYOMSQL;DATABASE=BorshDatabase;UID=sa;PWD=123')
    cursor = cnxn.cursor()
    cursor.execute(f'SELECT * FROM Users WHERE UserNumber = \'{Number}\'')
    for row in cursor:
        a = row[1]
        if Number == a:
            print("Вы успешно авторизовались!")
            return 1
    cnxn.close()

#Функция регистрации
def Registration(Number):
    cnxn = pyodbc.connect('DRIVER={SQL Server};SERVER=LAPTOP-BIF79HGK\ARTYOMSQL;DATABASE=BorshDatabase;UID=sa;PWD=123')
    cursor = cnxn.cursor()
    try:
        cursor.execute(f"INSERT INTO Users (UserNumber, Role_ID, UserBalance) VALUES (\'{Number}\', 1, \'0\')",)
        cnxn.commit()
    except:
        print("Вы неверно ввели номер!")
    cnxn.close()

#Функция проверки роли
def CheckRole(Number):
    cnxn = pyodbc.connect('DRIVER={SQL Server};SERVER=LAPTOP-BIF79HGK\ARTYOMSQL;DATABASE=BorshDatabase;UID=sa;PWD=123')
    cursor = cnxn.cursor()
    cursor.execute(f'SELECT * FROM Users WHERE UserNumber = \'{Number}\'')
    for row in cursor:
        a = row[2]
        if a == 2:
            return 2
        if a == 1:
            return 1
    cnxn.close()

#Рандомный баланс
def RandomBalance(Number):
    cnxn = pyodbc.connect('DRIVER={SQL Server};SERVER=LAPTOP-BIF79HGK\ARTYOMSQL;DATABASE=BorshDatabase;UID=sa;PWD=123')
    cursor = cnxn.cursor()
    cursor.execute(f'SELECT * FROM Users WHERE UserNumber = \'{Number}\'')
    for row in cursor:
        a = row[3]
        print(a)
        a += int(random.randint(350,800))
        print(a)
        cursor.execute(f"UPDATE Users set UserBalance =\'{a}\' where UserNumber = \'{Number}\'")
        cnxn.commit()
        cnxn.close()
        return int(a)
    
#Списание денег
def GetMoney(Balance, Number, Sum):
    cnxn = pyodbc.connect('DRIVER={SQL Server};SERVER=LAPTOP-BIF79HGK\ARTYOMSQL;DATABASE=BorshDatabase;UID=sa;PWD=123')
    cursor = cnxn.cursor()
    cursor.execute(f"UPDATE Users set UserBalance =\'{Balance}\' where UserNumber = \'{Number}\'")
    cnxn.commit()
    cnxn.close()
    try:
        cursor.execute(f'SELECT * FROM Users')
        for row in cursor:
            b = row[2]
            if b == 2:
                a = row[3]
                a += int(Sum)
                cursor.execute(f"UPDATE Users set UserBalance =\'{a}\' where Role_ID = \'{2}\'")
                cnxn.commit()
    except:
        print("")


#Сохранение истории
def SaveHistory(Number, Summ, cabbage, carrot, onion, parsley, dill, beetroot, meat, potato, garlic, tomatopaste, vinegar, beans):
    cnxn = pyodbc.connect('DRIVER={SQL Server};SERVER=LAPTOP-BIF79HGK\ARTYOMSQL;DATABASE=BorshDatabase;UID=sa;PWD=123')
    cursor = cnxn.cursor()
    try:
        cursor.execute(f'SELECT * FROM Users WHERE UserNumber = \'{Number}\'')
        for row in cursor:
            a = int(row[0])
            cursor.execute(f"INSERT INTO History (User_ID, Date, Summ, Cabbage, Carrot, Onion, Parsley, Dill, Beetroot, Meat, Potato, Garlic, TomatoPaste, Vinegar, Beans) VALUES (\'{a}\', \'{datetime.date.today()}\', \'{Summ}\', \'{cabbage}\', \'{carrot}\', \'{onion}\', \'{parsley}\', \'{dill}\', \'{beetroot}\', \'{meat}\', \'{potato}\', \'{garlic}\', \'{tomatopaste}\', \'{vinegar}\', \'{beans}\')",)
            cnxn.commit()
    except:
        print("")



#Проверка наличия продуктов
def CheckProduct(Name, Count):
    cnxn = pyodbc.connect('DRIVER={SQL Server};SERVER=LAPTOP-BIF79HGK\ARTYOMSQL;DATABASE=BorshDatabase;UID=sa;PWD=123')
    cursor = cnxn.cursor()
    cursor.execute(f'SELECT * FROM Ingredients WHERE Ingredient_Name = \'{Name}\'')
    for row in cursor:
        a = int(row[2])
        if (a==0):
            print("На складе нет данного продукта!")
            return 2
        if ((a - Count)<0 and a!=0):
            print("У нас нет столько продуктов!")
            return 0
        if ((a - Count)>=0):
            b = a - Count
            cursor.execute(f"UPDATE Ingredients set Quantity =\'{b}\' where Ingredient_Name = \'{Name}\'")
            cnxn.commit()
            cnxn.close()
            return 1
    cnxn.close()

#Пополнение запасов
def AdminPanel(Ingredient, Balance):
    cnxn = pyodbc.connect('DRIVER={SQL Server};SERVER=LAPTOP-BIF79HGK\ARTYOMSQL;DATABASE=BorshDatabase;UID=sa;PWD=123')
    cursor = cnxn.cursor()
    try:
        cursor.execute(f'SELECT * FROM Ingredients WHERE Ingredient_Name = \'{Ingredient}\'')
        for row in cursor:
            quantity = row[2]
            cost = row[3]
            print(f"Сколько вы хотите купить? 1 штука стоит {cost}")
            cout = int(input())
            count = cout+quantity
            blanc = int(Balance - cost*cout)
            cursor.execute(f"UPDATE Ingredients set Quantity =\'{count}\' where Ingredient_Name = \'{Ingredient}\'")
            cnxn.commit()
            GotMoneyAdmin(blanc)
    except:
        print("")



#Количество продуктов на данный момент
def AllProducts():
    cnxn = pyodbc.connect('DRIVER={SQL Server};SERVER=LAPTOP-BIF79HGK\ARTYOMSQL;DATABASE=BorshDatabase;UID=sa;PWD=123')
    cursor = cnxn.cursor()
    cursor.execute(f'SELECT * FROM Ingredients')
    print("Список продуктов на складе: ")
    for row in cursor:
        name = row[1]
        quantity = row[2]
        cost = row[3]
        print(f"Название: {name}| Количество: {quantity}| Цена: {cost}")
    cnxn.close()

def GotMoneyAdmin(Balance):
    cnxn = pyodbc.connect('DRIVER={SQL Server};SERVER=LAPTOP-BIF79HGK\ARTYOMSQL;DATABASE=BorshDatabase;UID=sa;PWD=123')
    cursor = cnxn.cursor()
    cursor.execute(f"UPDATE Users set UserBalance =\'{Balance}\' where Role_ID = \'{2}\'")
    cnxn.commit()
    cnxn.close()

def getHistory(Number):
    cnxn = pyodbc.connect('DRIVER={SQL Server};SERVER=LAPTOP-BIF79HGK\ARTYOMSQL;DATABASE=BorshDatabase;UID=sa;PWD=123')
    cursor = cnxn.cursor()
    cursor.execute(f'SELECT * FROM Users WHERE UserNumber = \'{Number}\'')
    for row in cursor:
        a = row[0]
        getByIDHistory(a)
    cnxn.close()


def getByIDHistory(ID):
    cnxn = pyodbc.connect('DRIVER={SQL Server};SERVER=LAPTOP-BIF79HGK\ARTYOMSQL;DATABASE=BorshDatabase;UID=sa;PWD=123')
    cursor = cnxn.cursor()
    cursor.execute(f'SELECT * FROM History where User_ID = \'{ID}\'')
    print("")
    for row in cursor:
        print(f'ID: {row[1]}, Дата: {row[2]}, Сумма: {row[3]}, Капуста: {row[4]}, Морковь: {row[5]}, Лук: {row[6]}, Петрушка: {row[7]}, Укроп: {row[8]}, Свекла: {row[9]}, Мясо: {row[10]}, Картошка: {row[11]}, Чеснок: {row[12]}, Томатная паста: {row[13]}, Уксус: {row[14]}, Фасоль: {row[15]}')
    print("")
    cnxn.close()

def getByIDCheck(Number, SUM, cok):
    cnxn = pyodbc.connect('DRIVER={SQL Server};SERVER=LAPTOP-BIF79HGK\ARTYOMSQL;DATABASE=BorshDatabase;UID=sa;PWD=123')
    cursor = cnxn.cursor()
    cursor.execute(f'SELECT * FROM Users WHERE UserNumber = \'{Number}\'')
    for row in cursor:
        a = row[0]
        CheckSave(Number, a, SUM, cok)
    cnxn.close()

def CheckSave(Number, ID, Sum, cockroach):
    cnxn = pyodbc.connect('DRIVER={SQL Server};SERVER=LAPTOP-BIF79HGK\ARTYOMSQL;DATABASE=BorshDatabase;UID=sa;PWD=123')
    cursor = cnxn.cursor()
    doc = docx.Document()
    doc.add_heading(f'Чек пользователя {Number}',0)
    cursor.execute(f'SELECT * FROM History where User_ID = \'{ID}\'')
    doc.add_paragraph(f'{Sum}')
    if (cockroach == 1):
       doc.add_paragraph('В еде попался таракан')
    doc.add_paragraph('Борщ')
    doc.add_paragraph(f'{datetime.datetime.now()}')
    doc.save(f'check{Sum}{datetime.datetime.now().hour+datetime.datetime.now().minute+datetime.datetime.now().microsecond}.docx')


cardType = int(0)

def LoyalCard(Number):
    cnxn = pyodbc.connect('DRIVER={SQL Server};SERVER=LAPTOP-BIF79HGK\ARTYOMSQL;DATABASE=BorshDatabase;UID=sa;PWD=123')
    cursor = cnxn.cursor()
    cursor.execute(f'SELECT * FROM Users WHERE UserNumber = \'{Number}\'')
    for row in cursor:
        a = row[0]
    cursor.close
    cnxn = pyodbc.connect('DRIVER={SQL Server};SERVER=LAPTOP-BIF79HGK\ARTYOMSQL;DATABASE=BorshDatabase;UID=sa;PWD=123')
    cursor = cnxn.cursor()
    cursor.execute(f'SELECT * FROM History where User_ID = \'{a}\'')
    b=int(0)
    for row in cursor:
        b += row[3]
    if (b >=5000 and b<15000):
            return 1 #Бронзовая карта
    if (b >= 15000 and b<25000):
        return 2 #Серебрянная карта
    if (b >25000):
        return 3 #Золотая карта
    if (b < 5000):
        return 0
    cursor.close


while True:
    while (Autoriz==0):
        print("Авторизоваться (1)  |  Зарегестрироваться (2)")
        choose = int(input())
        if (choose == 1):
            print("Введите свой номер телефона в формате +7(|||)|||-||-||")
            Number = input()
            Autoriz = Autorization(Number)
            if Autoriz == 1:
                Role = int(CheckRole(Number))
                Balance = int(RandomBalance(Number))
                cardType = int(LoyalCard(Number))
                if (cardType == 1):
                    print("Бронзовая карта лояльности 5%")
                if (cardType == 2):
                    print("Серебрянная карта лояльности 10%")
                if (cardType == 3):
                    print("Золотая карта лояльности 20%")
        if (choose == 2):
            print("Введите свой номер телефона в формате +7(|||)|||-||-||")
            Number = input()
            Registration(Number)
    if Role == 1:
        print("Стандартный борщ стоит 350 рублей. При увеличении количества продуктов в борще, у него увеличивается стоимость!")
        print("АКЦИЯ! ПЯТОЕ БЛЮДО В ПОДАРОК!")
        print("Выйти из магазина (1), купить (2) или посмотреть историю (3)? Выйти нельзя, пока не будет куплен хотя бы один товар!\n")
        ExitYes = int(input())
        if ExitYes == 3:
            getHistory(Number)
        if ExitYes == 1:
            if Challenge >= 1:
                break;
        if ExitYes == 2:
            print("Сколько порций борща вы желаете купить?\n")
            Quantity = int(input("Введите количество порций: "))
            if (Quantity < 0):
                print("Вы не можете взять отрицательное количество борщей!")
            if (Quantity == 0):
                print("Вы не можете взять 0 борщей. Нужен хотя бы один борщ!")    
            if (Quantity > 0):
                #Начало основной программы
                while Counter < Quantity:
                    Counter +=1
                    print("Вы желаете купить обычный борщ или собрать свой?")
                    Selecter = int(input("Свой - 1| Стандартный борщ - 2| : "))
                    if ((Balance-350)<0):
                        print("У вас не хватает средств")
                    else:
                        if (Selecter == 2):
                            SummCheck=350
                            count = int(0)
                            Summ+=SummCheck
                            j=int(CheckProduct("Свекла", 1))
                            j=int(CheckProduct("Мясо", 1))
                            j=int(CheckProduct("Картофель", 1))
                            j=int(CheckProduct("Луковица", 1))
                            j=int(CheckProduct("Морковь", 1))
                            j=int(CheckProduct("Капуста", 1))
                            j=int(CheckProduct("Чеснок", 1))
                            j=int(CheckProduct("Томатная паста", 1))
                            j=int(CheckProduct("Уксус", 1))
                            j=int(CheckProduct("Петрушка", 1))
                            j=int(CheckProduct("Укроп", 1))
                            if (cardType == 1):
                                count = SummCheck-SummCheck*0.05
                                SummCheck = count
                            if (cardType == 2):
                                count = SummCheck-SummCheck*0.1
                                SummCheck = count
                            if (cardType == 3):
                                count = SummCheck-SummCheck*0.2
                                SummCheck = count
                            if (j==1):
                                Challenge += 1
                                Cockroach = int(random.randint(1,6))
                                UserSees = int(random.randint(1,7))
                                if (Cockroach==5 and UserSees==5):
                                    print("ВЫ НАШЛИ ТАРАКАНА! ДЖЕКПОТ!")
                                    Summ-=SummCheck*0.3
                                    print(f"Скидка на данное блюдо составляет {SummCheck*0.3}. Цена: {SummCheck-SummCheck*0.3}")
                                    Balance -= int(SummCheck-SummCheck*0.3)
                                    count=SummCheck-SummCheck*0.3
                                    SummCheck=count
                                    cok = 1
                                else:
                                    Balance -= int(SummCheck)
                            SaveHistory(Number, SummCheck, 1, 1, 1, 1, 1, 1, 1, 1, 1, 1, 1, 0)
                            if (j==2 or j==0):
                                Summ -=350
                                Counter-=1
                        if (Selecter == 1):
                            SummCheck =350
                            print("Сколько листьев капусты вам порезать в борщ?")
                            Cabbage = int(input())
                            SummCheck += int(CabbageFunct(Cabbage))
                            print("Сколько морковок вам порезать в борщ?")
                            Carrot = int(input())
                            SummCheck += int(CarrotFunct(Carrot))
                            print("Сколько луковиц вам порезать в борщ?")
                            Union = int(input())
                            SummCheck += int(UnionFunct(Union))
                            print("Сколько листьев петрушки добавить в борщ?")
                            Parsley = int(input())
                            SummCheck += int(ParsleyFunct(Parsley))
                            print("Сколько листьев укропа добавить в борщ?")
                            Dill = int(input())
                            SummCheck += int(DillFunct(Dill))
                            print("Сколько свеклы добавить в борщ?")
                            Beetroot = int(input()) 
                            SummCheck += int(BeetrootFunct(Beetroot))
                            print("Сколько кусков мяса добавить в борщ? 1 кусок - это 450 грамм")
                            Meat = int(input())
                            SummCheck += int(MeatFunct(Meat))
                            print("Сколько картошек добавить в борщ?")
                            Potato = int(input())
                            SummCheck += int(PotatoFunct(Potato))
                            print("Сколько зубчиков чеснока добавить в борщ?")
                            Garlic = int(input())
                            SummCheck += int(GarlicFunct(Garlic))
                            print("Сколько ч. л. томатной пасты добавить в борщ?")
                            Tomato = int(input())
                            SummCheck += int(TomatoPasteFunct(Tomato))
                            print("Сколько ст. л. уксуса добавить в борщ?")
                            Vinegar = int(input())
                            SummCheck += int(VinegarFunct(Vinegar))
                            print("Желаете ли добавить фасоль из топинга в борщ? 1 - это одна упаковка 200 грамм.")
                            Beans = int(input())
                            SummCheck += int(BeansFunct(Beans))

                            Balance -= int(SummCheck)
                            Summ+=SummCheck
                            print (SummCheck)
                            Challenge += 1
                            Cockroach += int(random.randint(1,6))
                            UserSees += int(random.randint(1,7))
                            if (Cockroach==5 and UserSees==5):
                                print("ВЫ НАШЛИ ТАРАКАНА! ДЖЕКПОТ!")
                                
                                Summ-=SummCheck*0.3
                                print(f"Скидка на данное блюдо составляет {SummCheck*0.3}. Цена ")
                            if (cardType == 1):
                                count = SummCheck-SummCheck*0.05
                                SummCheck = count
                            if (cardType == 2):
                                count = SummCheck-SummCheck*0.1
                                SummCheck = count
                            if (cardType == 3):
                                count = SummCheck-SummCheck*0.2
                                SummCheck = count
                            SaveHistory(Number, Summ, Cabbage, Carrot, Onion, Parsley, Dill, Beetroot, Meat, Potato, Garlic, TomatoPaste, Vinegar, Beans)
                        if (int(Counter%5)==int(0)):
                            print("Каждая пятая покупка - станартный борщ в подарок!")
                            j = int(CheckProduct("Свекла", 1))
                            j = int(CheckProduct("Мясо", 1))
                            j=int( CheckProduct("Картофель", 1))
                            j=int(CheckProduct("Луковица", 1))
                            j=int(CheckProduct("Морковь", 1))
                            j=int(CheckProduct("Капуста", 1))
                            j=int(CheckProduct("Чеснок", 1))
                            j=int(CheckProduct("Томатная паста", 1))
                            j=int(CheckProduct("Уксус", 1))
                            j=int(CheckProduct("Петрушка", 1))
                            j=int(CheckProduct("Укроп", 1))
                            SaveHistory(Number, 0, 1, 1, 1, 1, 1, 1, 1, 1, 1, 1, 1, 0)
            if (Balance>0):
                GetMoney(Balance, Number, Summ)
                getByIDCheck(Number, Summ, cok)
                Counter = 0
            else:
                print("У вас не хватает средств!")
                cnxn = pyodbc.connect('DRIVER={SQL Server};SERVER=LAPTOP-BIF79HGK\ARTYOMSQL;DATABASE=BorshDatabase;UID=sa;PWD=123')
                cursor = cnxn.cursor()
                cursor.execute(f'SELECT * FROM Users WHERE UserNumber = \'{Number}\'')
                for row in cursor:
                    a = row[3]
                    Balance = int(a)
                    print(f"Текущий баланс: {Balance}")
                    Challenge-=1
    if Role == 2:
        AllProducts()
        print("\n")
        print("Вы желаете выйти (1) или купить продукты на склад (2)? Чтобы просмотреть историю покупок пользователя введите цифру 3.")
        ExitYes = int(input())
        if (ExitYes==1):
            break;
        if (ExitYes==2):
            Ingredient = input("Введите название ингредиента из списка: ")
            if (Ingredient=="Свекла" or Ingredient=="Мясо" or Ingredient=="Картофель" or Ingredient=="Луковица" or Ingredient=="Морковь" or Ingredient=="Капуста" or Ingredient=="Чеснок" or Ingredient=="Томатная паста" or Ingredient=="Уксус" or Ingredient=="Петрушка" or Ingredient=="Укроп" or Ingredient=="Фасоль"):
                AdminPanel(Ingredient, Balance)
            else:
                print("Вы ввели неверное название!")
        if (ExitYes==3):
            getHistory(input("Введите номер пользователя: "))
